from fastapi import FastAPI, BackgroundTasks
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import yt_dlp
import cv2
import pytesseract
import os
import numpy as np
from docx import Document
import uuid
import shutil
from fastapi import Request, Form
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
import platform


# ───────────── CONFIG ───────────── #
if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = os.getenv(
        "TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    )
else:
    pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_CMD", "tesseract")
# ───────────── FASTAPI APP ───────────── #
app = FastAPI()
templates = Jinja2Templates(directory="templates")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # adjust for your frontend
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ───────────── STATE ───────────── #
logs = []
status = {"running": False, "completed": False}
generated_file = None

def log(msg: str):
    logs.append(msg)
    print(msg)

# ───────────── UTILS ───────────── #
def download_video(url, filename="video.mp4"):
    ydl_opts = {"outtmpl": filename}
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        ydl.download([url])
    return filename
app.mount("/static", StaticFiles(directory="static"), name="static")

def extract_text_from_video(video_path, doc_path="save.docx", change_threshold=30):
    cap = cv2.VideoCapture(video_path)
    texts = []
    doc = Document()
    prev_frame = None
    frame_count = 0

    while cap.isOpened():
        ret, frame = cap.read()
        if not ret:
            break

        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        if prev_frame is not None:
            diff = cv2.absdiff(prev_frame, gray)
            non_zero_count = np.count_nonzero(diff)

            if non_zero_count > change_threshold * 1000:
                text = pytesseract.image_to_string(frame, lang="tam+eng+hin")
                if text.strip() and (len(texts) == 0 or text.strip() != texts[-1]):
                    texts.append(text.strip())
                    doc.add_paragraph(text.strip())
                    doc.add_paragraph("─" * 40)
                    log(f"📸 Screen changed → Extracted text at frame {frame_count}")

        prev_frame = gray
        frame_count += 1

    cap.release()
    doc.save(doc_path)
    log("✅ Document saved.")
    return doc_path

# ───────────── BACKGROUND PIPELINE ───────────── #
def run_pipeline(url):
    global status, generated_file
    try:
        status = {"running": True, "completed": False}
        logs.clear()
        generated_file = None

        log("⚡ Process started...")
        video_file = download_video(url)
        log("⬇️ Video downloaded.")

        # unique file for queue
        output_file = f"output_{uuid.uuid4().hex}.docx"
        doc_path = extract_text_from_video(video_file, output_file)
        log("📝 Text extraction complete.")

        generated_file = doc_path
        log("📂 File ready for download.")

        status = {"running": False, "completed": True}
    except Exception as e:
        log(f"❌ Error: {e}")
        status = {"running": False, "completed": True}

# ───────────── API MODELS ───────────── #
class ProcessRequest(BaseModel):
    url: str

# ───────────── ROUTES ───────────── #
@app.get("/", response_class=HTMLResponse)
def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})
@app.post("/process")
def process(request: ProcessRequest, background_tasks: BackgroundTasks):
    if status["running"]:
        return {"status": "Already running"}
    background_tasks.add_task(run_pipeline, request.url)
    return {"status": "Processing started"}

@app.get("/logs")
def get_logs():
    return {"logs": logs, "status": status}

@app.get("/download")
def download_file():
    if generated_file and os.path.exists(generated_file):
        return FileResponse(generated_file, filename=os.path.basename(generated_file))
    return {"error": "No file available"}

@app.post("/reset")
def reset():
    global logs, status, generated_file
    logs.clear()
    status = {"running": False, "completed": False}
    if generated_file and os.path.exists(generated_file):
        os.remove(generated_file)
    generated_file = None
    return {"status": "Reset complete"}
